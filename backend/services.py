"""
Text extraction and GitHub integration services.
"""

import re
import requests
from pathlib import Path
from datetime import datetime
from typing import Tuple, Dict, Any, List
import logging

# PDF processing with PyMuPDF
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# DOCX processing with python-docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# OpenAI integration
try:
    import openai
    import os
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

logger = logging.getLogger(__name__)


class TextExtractionService:
    """Service for extracting text from various file formats."""
    
    @staticmethod
    def extract_pdf_text(file_path: str) -> Tuple[str, bool]:
        """Extract text from PDF file using PyMuPDF."""
        if not PYMUPDF_AVAILABLE:
            return "PDF processing not available - install PyMuPDF: pip install PyMuPDF", False
        
        try:
            text_parts = []
            with fitz.open(file_path) as doc:
                for page in doc:
                    page_text = page.get_text() # type: ignore
                    if page_text.strip():
                        text_parts.append(page_text.strip())
            
            full_text = "\n".join(text_parts)
            return full_text, bool(full_text.strip())
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return f"Error reading PDF: {str(e)}", False
    
    @staticmethod
    def extract_docx_text(file_path: str) -> Tuple[str, bool]:
        """Extract text from DOCX file including headers, footers, and tables."""
        if not DOCX_AVAILABLE:
            return "DOCX processing not available - install python-docx: pip install python-docx", False
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # Extract from headers and footers
            for section in doc.sections:
                # Headers
                header = section.header
                for para in header.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text.strip())
                
                # Footers
                footer = section.footer
                for para in footer.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text.strip())
            
            full_text = "\n".join(text_parts)
            
            if len(full_text.strip()) > 10:  # Minimum text requirement
                return full_text, True
            else:
                return "No extractable text found. File may contain images or unsupported formatting.", False
                
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return f"Error reading DOCX: {str(e)}", False
    
    @staticmethod
    def extract_txt_text(file_path: str) -> Tuple[str, bool]:
        """Extract text from TXT file with multiple encoding support."""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                return text.strip(), bool(text.strip())
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error reading TXT file: {e}")
                return f"Error reading TXT: {str(e)}", False
        
        return "Could not decode text file with any supported encoding", False
    
    @classmethod
    def extract_text(cls, file_path: str, file_extension: str) -> Tuple[str, bool]:
        """Extract text based on file extension."""
        extraction_functions = {
            '.pdf': cls.extract_pdf_text,
            '.docx': cls.extract_docx_text,
            '.doc': cls.extract_docx_text,  # Treat .doc as .docx
            '.txt': cls.extract_txt_text,
        }
        
        extraction_func = extraction_functions.get(file_extension.lower())
        if not extraction_func:
            return f"Unsupported file type: {file_extension}. Supported: PDF, DOCX, TXT", False
        
        return extraction_func(file_path)


class ResumeAnalysisService:
    """Service for analyzing resume content with AI."""
    
    @staticmethod
    def generate_summary(text: str, max_length: int = 500) -> str:
        """Generate AI-powered summary using OpenAI or fallback to basic summary."""
        if not text or len(text.strip()) == 0:
            return "No content available for summary"
        
        # Try OpenAI first
        if OPENAI_AVAILABLE and os.getenv('OPENAI_API_KEY'):
            try:
                client = openai.OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
                response = client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[{
                        "role": "user",
                        "content": f"Summarize this resume in 2-3 professional sentences highlighting key qualifications and experience:\n\n{text[:3000]}"
                    }],
                    max_tokens=150,
                    temperature=0.7
                )
                return response.choices[0].message.content.strip() # type: ignore
            except Exception as e:
                logger.error(f"OpenAI summary generation failed: {e}")
                # Fall back to basic summary
        
        # Basic fallback summary
        return ResumeAnalysisService._basic_summary(text, max_length)
    
    @staticmethod
    def _basic_summary(text: str, max_length: int) -> str:
        """Generate basic summary without AI."""
        cleaned_text = ' '.join(text.split())
        
        if len(cleaned_text) <= max_length:
            return cleaned_text
        
        # Try to break at sentence boundaries
        sentences = re.split(r'[.!?]+', cleaned_text)
        summary_parts = []
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            sentence_with_punct = sentence + '. '
            if current_length + len(sentence_with_punct) <= max_length - 3:
                summary_parts.append(sentence_with_punct)
                current_length += len(sentence_with_punct)
            else:
                break
        
        if summary_parts:
            summary = ''.join(summary_parts).strip()
            if len(summary) < len(cleaned_text):
                summary += "..."
            return summary
        else:
            return cleaned_text[:max_length-3] + "..."
    
    @staticmethod
    def extract_contact_info(text: str) -> Dict[str, Any]:
        """Extract contact information from resume text."""
        info = {"email": None, "phone": None, "name": None}
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b'
        email_matches = re.findall(email_pattern, text, re.IGNORECASE)
        if email_matches:
            info["email"] = email_matches[0]
        
        # Extract phone
        phone_patterns = [
            r'\+?1?[-.\s]?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
            r'\b([0-9]{3})[-.]([0-9]{3})[-.]([0-9]{4})\b'
        ]
        
        for pattern in phone_patterns:
            phone_matches = re.findall(pattern, text)
            if phone_matches:
                if len(phone_matches[0]) == 3:
                    info["phone"] = f"({phone_matches[0][0]}) {phone_matches[0][1]}-{phone_matches[0][2]}" # type: ignore
                break
        
        # Extract name (first line that looks like a name)
        lines = text.split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if line and len(line.split()) >= 2 and len(line) < 50:
                # Simple heuristic: 2+ words, not too long, not email/phone
                if '@' not in line and not re.search(r'\d{3}', line):
                    info["name"] = line # type: ignore
                    break
        
        return info
    
    @staticmethod
    def extract_skills(text: str) -> List[str]:
        """Extract technical skills from resume text."""
        skill_keywords = [
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby',
            'go', 'rust', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql',
            
            # Web Technologies
            'react', 'angular', 'vue.js', 'node.js', 'express', 'django', 'flask',
            'fastapi', 'spring', 'laravel', 'rails', 'html', 'css', 'sass', 'bootstrap',
            
            # Databases
            'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle', 'cassandra',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'terraform',
            'ansible', 'chef', 'puppet',
            
            # Tools & Others
            'git', 'linux', 'unix', 'bash', 'powershell', 'nginx', 'apache',
            'machine learning', 'data science', 'artificial intelligence', 'blockchain'
        ]
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in skill_keywords:
            if skill in text_lower:
                found_skills.append(skill.title())
        
        # Remove duplicates and limit to 15
        return list(dict.fromkeys(found_skills))[:15]


class GitHubService:
    """Service for GitHub API integration."""
    
    @staticmethod
    def get_user_profile(username: str) -> Tuple[Dict[str, Any], bool]:
        """Fetch GitHub user profile data."""
        try:
            url = f"https://api.github.com/users/{username}"
            headers = {'User-Agent': 'ShadowPixel-Resume-Analyzer'}
            
            # Add GitHub token if available
            github_token = os.getenv('GITHUB_TOKEN')
            if github_token:
                headers['Authorization'] = f'token {github_token}'
            
            response = requests.get(url, headers=headers, timeout=10)
            
            if response.status_code == 200:
                return response.json(), True
            else:
                error_msg = f"GitHub API returned {response.status_code}"
                if response.status_code == 403:
                    error_msg += " (Rate limited - consider adding GITHUB_TOKEN)"
                return {"error": error_msg}, False
                
        except requests.RequestException as e:
            logger.error(f"Error fetching GitHub profile for {username}: {e}")
            return {"error": str(e)}, False
    
    @staticmethod
    def get_user_repositories(username: str, limit: int = 10) -> Tuple[List[Dict], bool]:
        """Fetch user's public repositories."""
        try:
            url = f"https://api.github.com/users/{username}/repos"
            headers = {'User-Agent': 'ShadowPixel-Resume-Analyzer'}
            params = {"sort": "updated", "per_page": limit}
            
            # Add GitHub token if available
            github_token = os.getenv('GITHUB_TOKEN')
            if github_token:
                headers['Authorization'] = f'token {github_token}'
            
            response = requests.get(url, headers=headers, params=params, timeout=10)
            
            if response.status_code == 200:
                repos = response.json()
                simplified_repos = []
                for repo in repos:
                    simplified_repos.append({
                        "name": repo["name"],
                        "description": repo.get("description", ""),
                        "language": repo.get("language", ""),
                        "stars": repo["stargazers_count"],
                        "forks": repo["forks_count"],
                        "updated_at": repo["updated_at"],
                        "url": repo["html_url"]
                    })
                return simplified_repos, True
            else:
                return [], False
                
        except requests.RequestException as e:
            logger.error(f"Error fetching GitHub repos for {username}: {e}")
            return [], False
