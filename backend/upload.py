"""
FastAPI Resume Upload and Processing Router.

This module handles resume file uploads with text extraction, validation,
and basic information parsing for multiple file formats.
"""

from typing import Dict, Tuple, Optional, List, Any, Set
from fastapi import APIRouter, UploadFile, Form, HTTPException, File
from fastapi.responses import JSONResponse
import logging
import os
import uuid
import re
from pathlib import Path
from datetime import datetime

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF not available. PDF processing will be disabled.")

try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX processing will be disabled.")

# Set up logging
logger = logging.getLogger(__name__)

# Create router with proper configuration
router = APIRouter(
    prefix="/resume", 
    tags=["resume"],
    responses={
        400: {"description": "Validation error"},
        500: {"description": "Internal server error"}
    }
)

# Configuration constants
UPLOAD_DIR: str = "uploaded_files"
ALLOWED_EXTENSIONS: Set[str] = {'.pdf', '.doc', '.docx', '.txt'}
MAX_FILE_SIZE: int = 5 * 1024 * 1024  # 5MB
MAX_USERNAME_LENGTH: int = 39
MAX_FILENAME_LENGTH: int = 255
MAX_SUMMARY_LENGTH: int = 500

# Create upload directory
os.makedirs(UPLOAD_DIR, exist_ok=True)

def validate_github_username(username: str) -> bool:
    """
    Validate GitHub username format according to GitHub rules.
    
    Args:
        username: The username to validate
        
    Returns:
        bool: True if valid, False otherwise
    """
    if not username or len(username) > MAX_USERNAME_LENGTH:
        return False
    
    # Single character must be alphanumeric
    if len(username) == 1:
        return username.isalnum()
    
    # Multi-character validation
    if username.startswith('-') or username.endswith('-') or '--' in username:
        return False
    
    # Only alphanumeric and hyphens allowed
    pattern = r'^[a-zA-Z0-9][a-zA-Z0-9-]*[a-zA-Z0-9]$'
    return bool(re.match(pattern, username))

def validate_file_security(filename: str) -> bool:
    """
    Validate filename for security issues.
    
    Args:
        filename: Filename to validate
        
    Returns:
        bool: True if safe, False if dangerous
    """
    if not filename:
        return False
    
    # Check for path traversal attempts and dangerous patterns
    dangerous_patterns = ['..', '/', '\\', '~', '<', '>', ':', '"', '|', '?', '*']
    return not any(pattern in filename for pattern in dangerous_patterns)

def validate_file(file: UploadFile) -> Tuple[bool, str]:
    """
    Validate uploaded file properties.
    
    Args:
        file: The uploaded file object
        
    Returns:
        Tuple[bool, str]: (is_valid, error_message)
    """
    if not file.filename:
        return False, "No filename provided"
    
    # Security check
    if not validate_file_security(file.filename):
        return False, "Invalid filename: contains dangerous characters"
    
    # Check filename length
    if len(file.filename) > MAX_FILENAME_LENGTH:
        return False, f"Filename too long (max {MAX_FILENAME_LENGTH} characters)"
    
    # Check file extension
    file_extension = Path(file.filename).suffix.lower()
    if file_extension not in ALLOWED_EXTENSIONS:
        return False, f"File type not allowed. Supported types: {', '.join(ALLOWED_EXTENSIONS)}"
    
    # Check if file extension has corresponding processor
    if file_extension == '.pdf' and not PYMUPDF_AVAILABLE:
        return False, "PDF processing not available (PyMuPDF not installed)"
    
    if file_extension == '.docx' and not DOCX_AVAILABLE:
        return False, "DOCX processing not available (python-docx not installed)"
    
    return True, "Valid"

def extract_pdf_text(file_path: str) -> Tuple[str, bool]:
    """
    Extract text from PDF file using PyMuPDF.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Tuple[str, bool]: (extracted_text, success_flag)
    """
    if not PYMUPDF_AVAILABLE:
        return "PDF processing not available (PyMuPDF not installed)", False
    
    try:
        text_parts = []
        with fitz.open(file_path) as doc:
            if len(doc) == 0:
                return "PDF file contains no pages", False
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()  # type: ignore[attr-defined]
                
                if page_text.strip():  # Only add non-empty pages
                    text_parts.append(f"--- Page {page_num + 1} ---")
                    text_parts.append(page_text.strip())
        
        if not text_parts:
            return "No text content found in PDF", False
        
        full_text = "\n".join(text_parts)
        return full_text, True
        
    except Exception as e:
        logger.error(f"Error extracting PDF text from {file_path}: {e}")
        return f"Error reading PDF: {str(e)}", False

def extract_docx_text(file_path: str) -> Tuple[str, bool]:
    """
    Extract text from DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Tuple[str, bool]: (extracted_text, success_flag)
    """
    if not DOCX_AVAILABLE:
        return "DOCX processing not available (python-docx not installed)", False
    
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if not text_parts:
            return "No text content found in DOCX", False
        
        full_text = "\n".join(text_parts)
        return full_text, True
        
    except Exception as e:
        logger.error(f"Error extracting DOCX text from {file_path}: {e}")
        return f"Error reading DOCX: {str(e)}", False

def extract_txt_text(file_path: str) -> Tuple[str, bool]:
    """
    Extract text from TXT file with encoding detection.
    
    Args:
        file_path: Path to the TXT file
        
    Returns:
        Tuple[str, bool]: (extracted_text, success_flag)
    """
    # Try different encodings in order of preference
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
            
            if not text.strip():
                return "Empty text file", False
            
            return text.strip(), True
            
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            return f"Error reading TXT file: {str(e)}", False
    
    return "Could not decode text file with any supported encoding", False

def extract_doc_text(file_path: str) -> Tuple[str, bool]:
    """
    Extract text from DOC file (legacy Word format).
    
    Args:
        file_path: Path to the DOC file
        
    Returns:
        Tuple[str, bool]: (extracted_text, success_flag)
    """
    # DOC files require specialized libraries like antiword or textract
    # For now, return a message indicating limitation
    return (
        "DOC file processing requires additional libraries (antiword, textract, or LibreOffice). "
        "Please convert to DOCX, PDF, or TXT format for processing.",
        False
    )

def extract_text_from_file(file_path: str, file_extension: str) -> Tuple[str, bool]:
    """
    Extract text based on file extension using appropriate processor.
    
    Args:
        file_path: Path to the file
        file_extension: File extension (e.g., '.pdf')
        
    Returns:
        Tuple[str, bool]: (extracted_text, success_flag)
    """
    extraction_functions = {
        '.pdf': extract_pdf_text,
        '.docx': extract_docx_text,
        '.txt': extract_txt_text,
        '.doc': extract_doc_text
    }
    
    extraction_func = extraction_functions.get(file_extension.lower())
    if not extraction_func:
        return f"Unsupported file type: {file_extension}", False
    
    return extraction_func(file_path)

def generate_summary(text: str, max_length: int = MAX_SUMMARY_LENGTH) -> str:
    """
    Generate a summary from extracted text with intelligent truncation.
    
    Args:
        text: Source text
        max_length: Maximum summary length
        
    Returns:
        str: Generated summary
    """
    if not text or len(text.strip()) == 0:
        return "No content available for summary"
    
    # Clean up the text
    cleaned_text = ' '.join(text.split())
    
    if len(cleaned_text) <= max_length:
        return cleaned_text
    
    # Try to break at sentence boundaries
    sentences = re.split(r'[.!?]+', cleaned_text)
    summary_parts = []
    current_length = 0
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
            
        # Check if adding this sentence would exceed limit
        sentence_with_punct = sentence + '. '
        if current_length + len(sentence_with_punct) <= max_length - 3:  # Reserve space for "..."
            summary_parts.append(sentence_with_punct)
            current_length += len(sentence_with_punct)
        else:
            break
    
    if summary_parts:
        summary = ''.join(summary_parts).strip()
        if len(summary) < len(cleaned_text):
            summary += "..."
        return summary
    else:
        # If no complete sentences fit, truncate at word boundary
        words = cleaned_text.split()
        truncated_words = []
        current_length = 0
        
        for word in words:
            if current_length + len(word) + 1 <= max_length - 3:
                truncated_words.append(word)
                current_length += len(word) + 1  # +1 for space
            else:
                break
        
        if truncated_words:
            return ' '.join(truncated_words) + "..."
        else:
            return cleaned_text[:max_length-3] + "..."

def extract_basic_info(text: str) -> Dict[str, Any]:
    """
    Extract basic information from resume text using pattern matching.
    
    Args:
        text: Resume text
        
    Returns:
        Dict[str, Any]: Dictionary containing extracted information
    """
    info: Dict[str, Any] = {
        "email": None,
        "phone": None,
        "skills": [],
        "years_experience": None,
        "education": []
    }
    
    if not text:
        return info
    
    # Extract email addresses
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b'
    email_matches = re.findall(email_pattern, text, re.IGNORECASE)
    if email_matches:
        # Filter out common non-email patterns
        valid_emails = [email for email in email_matches 
                       if not email.lower().endswith(('.png', '.jpg', '.gif', '.pdf'))]
        if valid_emails:
            info["email"] = valid_emails[0]  # Take first valid email
    
    # Extract phone numbers (various formats)
    phone_patterns = [
        r'\+?1?[-.\s]?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',  # US format
        r'\+([0-9]{1,3})[-.\s]?([0-9]{3,4})[-.\s]?([0-9]{3,4})[-.\s]?([0-9]{3,4})',  # International
        r'\b([0-9]{3})[-.]([0-9]{3})[-.]([0-9]{4})\b'  # Simple format
    ]
    
    for pattern in phone_patterns:
        phone_matches = re.findall(pattern, text)
        if phone_matches:
            # Reconstruct phone number from groups
            if len(phone_matches[0]) == 3:  # US format
                info["phone"] = f"({phone_matches[0][0]}) {phone_matches[0][1]}-{phone_matches[0][2]}"
            else:
                info["phone"] = '-'.join(phone_matches[0])
            break
    
    # Extract technical skills (expanded list)
    skill_keywords = [
        # Programming languages
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust',
        'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql',
        # Web technologies
        'react', 'angular', 'vue.js', 'node.js', 'express', 'django', 'flask', 'fastapi',
        'html', 'css', 'sass', 'less', 'bootstrap', 'tailwind',
        # Databases
        'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle', 'cassandra',
        # Cloud & DevOps
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'linux',
        # Data & ML
        'machine learning', 'deep learning', 'artificial intelligence', 'data science',
        'pandas', 'numpy', 'tensorflow', 'pytorch', 'scikit-learn',
        # Other
        'rest api', 'graphql', 'microservices', 'agile', 'scrum'
    ]
    
    text_lower = text.lower()
    found_skills = []
    for skill in skill_keywords:
        if skill in text_lower:
            found_skills.append(skill.title())
    
    # Remove duplicates and limit
    info["skills"] = list(dict.fromkeys(found_skills))[:15]
    
    # Extract years of experience
    experience_patterns = [
        r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
        r'experience[:\s]*(\d+)\+?\s*years?',
        r'(\d+)\+?\s*years?\s*in\s*\w+'
    ]
    
    for pattern in experience_patterns:
        matches = re.findall(pattern, text_lower)
        if matches:
            try:
                years = int(matches[0])
                if 0 <= years <= 50:  # Reasonable range
                    info["years_experience"] = years
                    break
            except (ValueError, IndexError):
                continue
    
    # Extract education information
    education_keywords = [
        'bachelor', 'master', 'phd', 'doctorate', 'mba', 'bs', 'ms', 'ba', 'ma',
        'computer science', 'engineering', 'mathematics', 'physics', 'chemistry',
        'business administration', 'information technology'
    ]
    
    found_education = []
    for keyword in education_keywords:
        if keyword in text_lower:
            found_education.append(keyword.title())
    
    info["education"] = list(dict.fromkeys(found_education))[:5]
    
    return info

@router.post("/upload/")
async def upload_resume(
    resume: UploadFile = File(..., description="Resume file (PDF, DOCX, TXT)"),
    github: str = Form(..., description="GitHub username", min_length=1, max_length=39)
) -> JSONResponse:
    """
    Upload and process a resume file with comprehensive text extraction.
    
    Supports PDF, DOCX, and TXT files with automatic text extraction,
    summary generation, and basic information parsing.
    
    Args:
        resume: Uploaded resume file
        github: GitHub username for file organization
        
    Returns:
        JSONResponse containing processing results and extracted information
        
    Raises:
        HTTPException: For validation errors or processing failures
    """
    file_location: Optional[str] = None
    
    try:
        # Validate GitHub username
        if not validate_github_username(github):
            raise HTTPException(
                status_code=400,
                detail=(
                    "Invalid GitHub username format. Must be 1-39 characters, "
                    "alphanumeric and hyphens only, cannot start/end with hyphen or contain consecutive hyphens."
                )
            )
        
        # Validate file
        is_valid, error_message = validate_file(resume)
        if not is_valid:
            raise HTTPException(status_code=400, detail=error_message)
        
        # Read file contents
        contents = await resume.read()
        
        # Validate file size
        if len(contents) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File too large. Maximum size: {MAX_FILE_SIZE // (1024*1024)}MB. "
                       f"Current file size: {len(contents) / (1024*1024):.2f}MB"
            )
        
        # Generate secure, unique filename
        if resume.filename is None:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        file_extension = Path(resume.filename).suffix.lower()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{github}_{timestamp}_{uuid.uuid4().hex[:8]}{file_extension}"
        file_location = os.path.join(UPLOAD_DIR, unique_filename)
        
        # Save file securely
        try:
            with open(file_location, "wb") as f:
                f.write(contents)
        except IOError as e:
            logger.error(f"Failed to save file {unique_filename}: {e}")
            raise HTTPException(status_code=500, detail="Failed to save uploaded file")
        
        # Extract text from file
        extracted_text, extraction_success = extract_text_from_file(file_location, file_extension)
        
        # Generate summary and extract information only if extraction succeeded
        if extraction_success:
            summary = generate_summary(extracted_text)
            basic_info = extract_basic_info(extracted_text)
            text_length = len(extracted_text)
        else:
            summary = extracted_text  # Error message
            basic_info = {}
            text_length = 0
        
        # Prepare comprehensive response
        response_data = {
            "message": "Resume uploaded and processed successfully!" if extraction_success 
                      else "Resume uploaded but text extraction failed",
            "github_username": github,
            "original_filename": resume.filename,
            "saved_filename": unique_filename,
            "file_size": len(contents),
            "file_size_human": f"{len(contents) / (1024*1024):.2f} MB",
            "file_type": file_extension,
            "extraction_success": extraction_success,
            "summary": summary,
            "extracted_info": basic_info,
            "text_length": text_length,
            "processing_timestamp": datetime.now().isoformat(),
            "capabilities": {
                "pdf_processing": PYMUPDF_AVAILABLE,
                "docx_processing": DOCX_AVAILABLE
            }
        }
        
        # Log processing results
        status = "Success" if extraction_success else "Failed"
        logger.info(
            f"✅ Processed resume: {resume.filename} -> {unique_filename} "
            f"for {github}. Extraction: {status}. Text length: {text_length}"
        )
        
        return JSONResponse(response_data, status_code=200)
        
    except HTTPException:
        # Clean up file if processing failed and file was created
        if file_location and os.path.exists(file_location):
            try:
                os.remove(file_location)
            except OSError:
                logger.warning(f"Could not clean up file: {file_location}")
        raise
        
    except Exception as e:
        # Clean up file if processing failed and file was created
        if file_location and os.path.exists(file_location):
            try:
                os.remove(file_location)
            except OSError:
                logger.warning(f"Could not clean up file: {file_location}")
        
        logger.error(f"❌ Unexpected error processing resume: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Internal server error during file processing"
        )

@router.get("/health", response_model=Dict[str, Any])
async def health_check() -> Dict[str, Any]:
    """
    Health check endpoint with capability information.
    
    Returns:
        Dict containing service status and available capabilities
    """
    return {
        "status": "healthy",
        "service": "Resume Upload and Processing Service",
        "version": "2.0.0",
        "capabilities": {
            "pdf_processing": PYMUPDF_AVAILABLE,
            "docx_processing": DOCX_AVAILABLE,
            "txt_processing": True,
            "doc_processing": False
        },
        "limits": {
            "max_file_size_mb": MAX_FILE_SIZE // (1024 * 1024),
            "max_username_length": MAX_USERNAME_LENGTH,
            "max_summary_length": MAX_SUMMARY_LENGTH
        },
        "timestamp": datetime.now().isoformat()
    }

@router.get("/supported-formats", response_model=Dict[str, Any])
async def get_supported_formats() -> Dict[str, Any]:
    """
    Get comprehensive information about supported file formats and features.
    
    Returns:
        Dict containing supported formats and feature details
    """
    format_details = {
        ".pdf": {
            "name": "Portable Document Format",
            "supported": PYMUPDF_AVAILABLE,
            "features": ["text_extraction", "multi_page_support"] if PYMUPDF_AVAILABLE else [],
            "note": "Requires PyMuPDF library" if not PYMUPDF_AVAILABLE else None
        },
        ".docx": {
            "name": "Microsoft Word Document",
            "supported": DOCX_AVAILABLE,
            "features": ["text_extraction", "table_extraction"] if DOCX_AVAILABLE else [],
            "note": "Requires python-docx library" if not DOCX_AVAILABLE else None
        },
        ".txt": {
            "name": "Plain Text File",
            "supported": True,
            "features": ["text_extraction", "encoding_detection"],
            "note": None
        },
        ".doc": {
            "name": "Microsoft Word Document (Legacy)",
            "supported": False,
            "features": [],
            "note": "Not supported. Please convert to DOCX or PDF format."
        }
    }
    
    return {
        "supported_formats": [fmt for fmt, details in format_details.items() if details["supported"]],
        "format_details": format_details,
        "max_file_size_mb": MAX_FILE_SIZE // (1024 * 1024),
        "processing_features": [
            "Text extraction",
            "Summary generation", 
            "Email extraction",
            "Phone number extraction", 
            "Skills identification",
            "Experience years detection",
            "Education information extraction",
            "File validation and security checks"
        ],
        "github_username_rules": [
            "1-39 characters long",
            "Alphanumeric characters and hyphens only",
            "Cannot start or end with hyphen",
            "Cannot contain consecutive hyphens"
        ]
    }

@router.get("/stats", response_model=Dict[str, Any])
async def get_processing_stats() -> Dict[str, Any]:
    """
    Get statistics about processed files.
    
    Returns:
        Dict containing file processing statistics
    """
    try:
        total_files = 0
        total_size = 0
        file_types: Dict[str, int] = {}
        
        if os.path.exists(UPLOAD_DIR):
            for filename in os.listdir(UPLOAD_DIR):
                file_path = os.path.join(UPLOAD_DIR, filename)
                if os.path.isfile(file_path):
                    total_files += 1
                    
                    # Get file size
                    try:
                        file_size = os.path.getsize(file_path)
                        total_size += file_size
                    except OSError:
                        continue
                    
                    # Count file types
                    ext = Path(filename).suffix.lower()
                    file_types[ext] = file_types.get(ext, 0) + 1
        
        return {
            "total_files_processed": total_files,
            "total_storage_used": total_size,
            "total_storage_human": f"{total_size / (1024*1024):.2f} MB",
            "file_types_processed": file_types,
            "average_file_size": f"{(total_size / total_files) / (1024*1024):.2f} MB" if total_files > 0 else "0 MB",
            "upload_directory": UPLOAD_DIR,
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        logger.error(f"Error getting processing stats: {e}")
        raise HTTPException(status_code=500, detail="Error retrieving statistics")
